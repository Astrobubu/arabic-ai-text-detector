from __future__ import annotations

import io
from pathlib import Path
from typing import Union


class ExtractionUnavailable(RuntimeError):
    """Raised when the optional 'docs' extra isn't installed for a given file type."""


class UnsupportedFileType(ValueError):
    """Raised for file extensions with no known extractor."""


def extract_text(data: Union[str, bytes, Path], filename: str | None = None) -> str:
    """Extract plain text from .txt/.docx/.pdf input.

    `data` may be a filesystem path, raw file bytes, or an already-decoded str
    (returned as-is). `filename` picks the extractor by extension and is
    required whenever `data` is bytes.
    """
    if isinstance(data, Path):
        filename = filename or data.name
        data = data.read_bytes()

    if isinstance(data, str):
        return data

    if filename is None:
        raise ValueError("filename is required to determine how to extract raw bytes.")

    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        return data.decode("utf-8", errors="replace")
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix == ".pdf":
        return _extract_pdf(data)
    raise UnsupportedFileType(f"Unsupported file type: {suffix or '(none)'}. Supported: .txt, .docx, .pdf")


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Reading .docx files needs the optional 'docs' extra. "
            'Install with: pip install -e ".[docs]"'
        ) from exc

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Reading .pdf files needs the optional 'docs' extra. "
            'Install with: pip install -e ".[docs]"'
        ) from exc

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)
